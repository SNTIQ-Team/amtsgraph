"""DOCX letter rendering from the org DIN-5008 template.

`templates/letter_base.docx` is the organization's 2nd_Antwort template
with the logo removed and the first-person RDG note in the page footer
(adapted once, see git history). At request time we only:

  - replace the %PLACEHOLDER% fields (paragraph-level, run-merging);
  - expand the body ('Mustertext' anchor) into real paragraphs;
  - fill or drop the optional lines (Per Fax vorab, Tel./Fax/E-Mail,
    Unser Az., Ihr Schreiben vom) — a label whose value is empty is
    removed rather than left dangling;
  - replace the Anlagen stub (Anlage X/Y/Z) or drop the whole block.

PDF conversion happens via a running unoserver (LibreOffice headless,
own systemd unit); callers fall back to the fpdf renderer when the
converter is unavailable. DOCX output needs no converter at all.
"""
from __future__ import annotations

import io
import os
import re
import subprocess
from pathlib import Path

import docx
from docx.text.paragraph import Paragraph

TEMPLATE = Path(os.environ.get(
    "SNTIQ_DOCX_TEMPLATE",
    Path(__file__).resolve().parent / "templates" / "letter_base.docx"))

UNOCONVERT = os.environ.get("SNTIQ_UNOCONVERT", "/usr/bin/unoconvert")


def _set_text(p: Paragraph, text: str) -> None:
    """Replace a paragraph's text, keeping the first run's formatting."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _delete(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)


def render_docx(
    *,
    sender: list[str],
    recipient: list[str],
    ort: str | None,
    date: str,
    fax_vorab: str | None,
    info_lines: list[str],
    subject: str,
    body: str,
    signature_name: str | None,
    attachments: str | None,
) -> bytes:
    doc = docx.Document(TEMPLATE)

    # ---- sender decomposition (name / street / "PLZ Ort") ----------
    s_name = sender[0] if sender else ""
    s_street = sender[1] if len(sender) > 1 else ""
    s_plzort = sender[2] if len(sender) > 2 else ""
    m = re.match(r"^\s*(\d{4,5})\s+(.*)$", s_plzort)
    s_plz, s_ort = (m.group(1), m.group(2)) if m else ("", s_plzort)
    if ort:
        s_ort = s_ort or ort

    r_lines = [l for l in recipient if l.strip()]
    info = [l for l in info_lines if l.strip()]
    mapping = {
        "%NAME_A%": s_name,
        "%STR_A%": s_street,
        "%POST_A%": s_plz,
        "%ORT_A%": s_ort or (ort or ""),
        "%NAME_B%": r_lines[0] if r_lines else "",
        "%ADR_1%": r_lines[1] if len(r_lines) > 1 else "",
        "%ADR_2%": r_lines[2] if len(r_lines) > 2 else "",
        "%ADR_3%": ", ".join(r_lines[3:]) if len(r_lines) > 3 else "",
        "%DATE%": date,
        "%FAX_B%": fax_vorab or "",
        "%TEL_A%": "",
        "%FAX_A%": "",
        "%EMAIL_A%": "",
        "%B_ACT%": info[0].split(":", 1)[-1].strip() if info else "–",
        "%B_DATE%": info[1].split(":", 1)[-1].strip() if len(info) > 1 else "–",
        "%S_AZ%": "–",
        "%S_DATE%": "–",
        "%THEME_1%": subject,
        "%DATE_B%": "",
    }

    def walk_paragraphs():
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    # ---- placeholder replacement -----------------------------------
    for p in list(walk_paragraphs()):
        if "%" not in p.text:
            continue
        new = p.text
        for k, v in mapping.items():
            new = new.replace(k, v)
        if new != p.text:
            _set_text(p, new)

    # ---- drop dangling optional lines -------------------------------
    # they live as soft-break LINES inside shared paragraphs (and carry
    # zero-width characters), so cleanup is line-based per paragraph
    ZW = dict.fromkeys(map(ord, "​‌‍﻿⁠"))
    LABEL = re.compile(
        r"^(Per Fax vorab:|Tel\.:|Fax:|E-Mail:|Ihr Schreiben vom)$")
    EMPTY_PAIR = re.compile(r"^[–-]\s*/\s*[–-]$")
    PAIR_LABEL = ("Unser Az. / Schr. v.", "Unser Az..", "Unser Az.")

    def _clean_lines(text: str) -> str:
        out: list[str] = []
        for ln in text.split("\n"):
            n = ln.translate(ZW).strip()
            if LABEL.match(n):
                continue
            if EMPTY_PAIR.match(n):
                if out and out[-1].translate(ZW).strip() in PAIR_LABEL:
                    out.pop()
                continue
            out.append(ln)
        # collapse runs of blank lines left behind
        collapsed: list[str] = []
        for ln in out:
            if ln.strip() == "" and collapsed and collapsed[-1].strip() == "":
                continue
            collapsed.append(ln)
        return "\n".join(collapsed)

    for p in list(walk_paragraphs()):
        cleaned = _clean_lines(p.text)
        if cleaned != p.text:
            _fill_simple(p, cleaned)

    # ---- body --------------------------------------------------------
    for p in list(doc.paragraphs):
        if p.text.strip() == "Mustertext":
            paras = body.split("\n\n")
            # drop the duplicated salutation/closing: template has its own
            if paras and paras[0].startswith("Sehr geehrte"):
                paras = paras[1:]
            if paras and paras[-1].strip() == "Mit freundlichen Grüßen":
                paras = paras[:-1]
            _fill_simple(p, paras[0] if paras else "")
            anchor = p
            for text in paras[1:]:
                np = _insert_after(anchor, doc)
                _fill_simple(np, text)
                spacer = _insert_after(np, doc)
                _move_before(spacer, np)
                _fill_simple(spacer, "")
                anchor = np
            break

    # ---- signature: %NAME_A% under the dotted line was already filled;
    # hand-fill mode blanks it so the user writes the name by pen
    if signature_name == "":
        # hand-fill mode: blank the name under the dotted line
        for p in doc.paragraphs:
            if p.text.strip() == s_name and "…" in (_prev_text(p) or ""):
                _set_text(p, "")

    # ---- Anlagen -------------------------------------------------------
    stub = [p for p in doc.paragraphs
            if p.text.strip() in ("Anlagen", "Anlage X", "Anlage Y",
                                  "Anlage Z")]
    if attachments:
        lines = [l.strip() for l in
                 attachments.replace("Anlagen:", "").split(",")]
        lines = [l for l in lines if l]
        keep = stub[0] if stub else None  # the "Anlagen" heading
        for p in stub[1:]:
            _delete(p)
        if keep is not None:
            anchor = keep
            for line in lines:
                np = _insert_after(anchor, doc)
                _fill_simple(np, line)
                anchor = np
    else:
        for p in stub:
            _delete(p)

    # ---- Rücksendeangabe: small bullet, keep 8pt ---------------------
    from docx.shared import Mm, Pt
    for p in walk_paragraphs():
        if "●" in p.text:
            _fill_simple(p, p.text.replace(" ● ", " • "))
            for r in p.runs:
                r.font.size = Pt(8)
            break

    # ---- Ort/Datum row: identical single tabs + explicit tab stops --
    cell0 = doc.tables[0].rows[1].cells[0]
    label_p = value_p = None
    for p in cell0.paragraphs:
        if p.text.replace("\t", " ").strip().startswith("Ort"):
            label_p = p
        elif label_p is not None and value_p is None and p.text.strip():
            value_p = p
    if label_p is not None and value_p is not None:
        vals = [v for v in re.split(r"\t+", value_p.text.strip()) if v]
        _fill_simple(label_p, "Ort\tDatum\tIhr Az. / Schr. v.")
        _fill_simple(value_p, "\t".join(vals[:3]))
        for p in (label_p, value_p):
            ts = p.paragraph_format.tab_stops
            for t in list(ts):
                pass
            ts.add_tab_stop(Mm(28))
            ts.add_tab_stop(Mm(58))

    # ---- signature: dots sized to the name; extra co-signers ---------
    names = [n for n in (signature_name or "").split("\n") if n.strip()]
    main = names[0] if names else ""
    dots_p = name_p = None
    for p in doc.paragraphs:
        if set(p.text.strip()) == {"…"}:
            dots_p = p
        elif dots_p is not None and name_p is None and (
                p.text.strip() == s_name or p.text.strip() == main
                or (signature_name == "" and p.text.strip() == "")):
            name_p = p
            break
    def _dots_for(n: str) -> str:
        return "…" * max(14, min(34, len(n) + 6))
    if dots_p is not None:
        _fill_simple(dots_p, _dots_for(main or s_name))
    if name_p is not None and names:
        _fill_simple(name_p, main)
        anchor = name_p
        for extra in names[1:]:
            gap = _insert_after(anchor, doc)
            _fill_simple(gap, "")
            dp = _insert_after(gap, doc)
            _fill_simple(dp, _dots_for(extra))
            npn = _insert_after(dp, doc)
            _fill_simple(npn, extra)
            anchor = npn

    # ---- drop the trailing empty-paragraph desert (blank page 2) -----
    body_paras = doc.paragraphs
    # walk from the end, removing empties; keep at most 1 after content
    tail_empty = 0
    for p in reversed(body_paras):
        if p.text.strip() == "":
            tail_empty += 1
            if tail_empty > 1:
                _delete(p)
        else:
            break
    # collapse long runs of empties anywhere after the signature
    run = []
    for p in list(doc.paragraphs):
        if p.text.strip() == "":
            run.append(p)
        else:
            if len(run) > 2:
                for extra_p in run[2:]:
                    _delete(extra_p)
            run = []
    if len(run) > 2:
        for extra_p in run[2:]:
            _delete(extra_p)

    # ---- DIN spacing: the subject starts ~3 blank lines below the
    # letterhead — the empty-paragraph collapse above is too eager there
    subj_first = subject.split("\n")[0][:40]
    for i, p in enumerate(doc.paragraphs):
        if subj_first and p.text.startswith(subj_first):
            blanks = 0
            j = i - 1
            while j >= 0 and doc.paragraphs[j].text.strip() == "":
                blanks += 1
                j -= 1
            for _ in range(max(0, 3 - blanks)):
                p.insert_paragraph_before("")
            break

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fill_simple(p: Paragraph, text: str) -> None:
    """Set paragraph text; '\n' becomes a soft line break."""
    for r in list(p.runs):
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run("")
    lines = text.split("\n")
    run.text = lines[0]
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _insert_after(p: Paragraph, doc) -> Paragraph:
    import copy as _copy
    new = doc.add_paragraph()
    p._element.addnext(new._element)
    new.style = p.style
    # carry direct paragraph formatting (justification, spacing …)
    if p._element.pPr is not None:
        new._element.insert(0, _copy.deepcopy(p._element.pPr))
    return new


def _move_before(p: Paragraph, ref: Paragraph) -> None:
    ref._element.addprevious(p._element)


def _prev_text(p: Paragraph) -> str | None:
    prev = p._element.getprevious()
    if prev is None:
        return None
    return "".join(prev.itertext())


def _add_marks(pdf_bytes: bytes) -> bytes:
    """Stamp print aids onto every page: DIN fold marks (105/210 mm),
    punch mark (148.5 mm) and the page number — drawn by us, because
    LibreOffice kept rendering NUMPAGES fields with rogue formatting."""
    try:
        from io import BytesIO
        from fpdf import FPDF
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(BytesIO(pdf_bytes))
        total = len(reader.pages)
        marks = FPDF("P", "mm", "A4")
        marks.set_auto_page_break(False)
        for i in range(total):
            marks.add_page()
            marks.set_draw_color(150, 150, 150)
            marks.set_line_width(0.25)
            marks.line(3, 105, 6, 105)      # Falzmarke 1
            marks.line(3, 148.5, 9, 148.5)  # Lochmarke (longer, centre)
            marks.line(3, 210, 6, 210)      # Falzmarke 2
            marks.set_font("helvetica", "", 6.5)
            marks.set_text_color(120, 120, 120)
            marks.set_xy(0, 289)
            marks.cell(210, 3, f"- Seite {i + 1} von {total} -",
                       align="C")
        overlays = PdfReader(BytesIO(bytes(marks.output()))).pages
        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            page.merge_page(overlays[i])
            writer.add_page(page)
        out = BytesIO()
        writer.write(out)
        return out.getvalue()
    except Exception:  # noqa: BLE001 — aids are a nicety, never fatal
        return pdf_bytes


def docx_to_pdf(docx_bytes: bytes, timeout: int = 40) -> bytes | None:
    """Convert via the local unoserver; None when unavailable/failed.

    Streams stdin→stdout ('-' '-'): the bytes travel over the RPC
    socket, so the api and unoserver sandboxes (PrivateTmp) need no
    shared filesystem.
    """
    if not Path(UNOCONVERT).exists():
        return None
    try:
        proc = subprocess.run(
            [UNOCONVERT, "--convert-to", "pdf", "-", "-"],
            input=docx_bytes, capture_output=True, timeout=timeout)
    except Exception:  # noqa: BLE001 — caller falls back to fpdf
        return None
    out = proc.stdout
    if proc.returncode != 0 or not out.startswith(b"%PDF"):
        return None
    return _add_marks(out)
